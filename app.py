import os
import re
import json
from datetime import datetime
from copy import deepcopy

import chainlit as cl
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document


load_dotenv()

MODEL_NAME = os.getenv("MODEL_NAME", "openrouter/free")
API_BASE = os.getenv("API_BASE", "https://openrouter.ai/api/v1")
API_KEY = os.getenv("API_KEY")

client = OpenAI(api_key=API_KEY, base_url=API_BASE)

DOMAIN_CATALOG_PATH = "domain_catalog.json"
GUARDRAIL_CHUNKS_PATH = "guardrail_chunks.json"
BRD_TEMPLATE_PATH = "data/brd_template.docx"
OUTPUT_DIR = "output"
AGENT_IMAGE_PATH = "public/agents.png"

os.makedirs(OUTPUT_DIR, exist_ok=True)

AGENT_ICONS = {
    "Scout": "🦅",
    "Inspector": "🦉",
    "Strategist": "🐺",
    "Scribe": "🐙",
    "Conductor": "🐬",
    "Guardian": "🦏",
}

CUSTOMER_JOURNEYS = {
    "1": {
        "name": "KVKK / PII Risky Request",
        "sample": "We need a dataset including customer phone number, TCKN, location and package information for campaign targeting."
    },
    "2": {
        "name": "Expired Retention Data Request",
        "sample": "We want to re-analyze campaign performance data from 5 years ago."
    },
    "3": {
        "name": "Existing Data Discovery and Routing",
        "sample": "We want to understand which platform contains the data required for the customer segmentation report."
    }
}

REQUIRED_FIELDS = {
    "business_need": "Business Need",
    "data_domain": "Data Domain",
    "source_system": "Source System",
    "expected_output": "Expected Output",
    "consumer": "Consumer",
    "frequency": "Frequency",
    "contains_pii": "Contains PII",
    "deadline": "Deadline",
}


BILINGUAL_TERMS = {
    "campaign": ["campaign", "campaigns", "marketing", "promotion", "kampanya", "pazarlama"],
    "marketing": ["marketing", "campaign", "kampanya", "pazarlama", "nps", "profitability"],
    "performance": ["performance", "kpi", "metric", "analytics", "performans", "metrik", "analitik"],
    "retention": ["retention", "historical", "archive", "5 years", "saklama", "veri saklama", "arşiv", "eski veri", "geçmiş veri"],
    "pii": ["pii", "personal data", "kvkk", "kişisel veri", "hassas veri", "tckn", "telefon", "msisdn", "location", "lokasyon"],
    "approval": ["approval", "approve", "review", "onay", "istisna onayı", "mimari kurul", "veri mimarisi kurulu"],
    "customer": ["customer", "subscriber", "musteri", "müşteri", "abone", "segment"],
    "source": ["source", "system", "platform", "kaynak", "sistem", "platform"],
}


DOMAIN_BOOST_RULES = {
    "Marketing": ["campaign", "campaigns", "marketing", "nps", "profitability", "promotion", "campaign performance"],
    "Customer": ["customer id", "customer identity", "customer type", "tariff", "customer"],
    "Product & Service": ["tariff", "bundle", "pricing", "product", "service"],
    "Customer Service Performance": ["aht", "fcr", "call center", "service efficiency"],
    "Management Reporting": ["executive", "market share", "arpu", "kpi"],
    "Network Capacity & Performance Mgmt": ["network performance", "capacity"],
    "Digital Channel Performance": ["digital analytics", "heatmap", "page load"],
}


def normalize(text):
    text = str(text or "").lower()
    text = text.replace("ı", "i").replace("ğ", "g").replace("ü", "u")
    text = text.replace("ş", "s").replace("ö", "o").replace("ç", "c")
    return re.sub(r"[^a-z0-9 ]", " ", text)


def is_empty(value):
    return value in [None, "", "-", "Unknown", "unknown", "N/A", "n/a", "Not found"]


def safe_json_loads(text):
    try:
        return json.loads(text)
    except Exception:
        match = re.search(r"\{.*\}", text or "", re.DOTALL)
        if match:
            return json.loads(match.group(0))
        return {}


def load_json(path):
    if not os.path.exists(path):
        return []
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def expand_query(text):
    expanded = normalize(text)
    for key, terms in BILINGUAL_TERMS.items():
        if key in expanded or any(normalize(t) in expanded for t in terms):
            expanded += " " + " ".join(terms)
    return normalize(expanded)


def score_domain(request, domain):
    request_n = expand_query(request)
    searchable = normalize(
        f"{domain.get('domain_name')} "
        f"{domain.get('category')} "
        f"{domain.get('definition')} "
        f"{domain.get('sample_data')} "
        f"{domain.get('data_owner')} "
        f"{domain.get('related_sub_team')}"
    )

    score = 0
    reasons = []

    for term in set(request_n.split()):
        if len(term) > 2 and term in searchable:
            score += 1
            reasons.append(f"keyword match: {term}")

    domain_name = domain.get("domain_name", "")

    for boost_term in DOMAIN_BOOST_RULES.get(domain_name, []):
        if normalize(boost_term) in request_n:
            score += 5
            reasons.append(f"domain boost: {boost_term}")

    if "campaign performance" in request_n and domain_name == "Marketing":
        score += 10
        reasons.append("strong phrase match: campaign performance")

    if "performance" in request_n and "performance" in searchable:
        score += 3
        reasons.append("performance match")

    return score, reasons


def scout_agent(user_request):
    domains = load_json(DOMAIN_CATALOG_PATH)

    scored = []
    for domain in domains:
        score, reasons = score_domain(user_request, domain)
        scored.append({
            "score": score,
            "domain": domain,
            "reasons": reasons[:5]
        })

    scored.sort(key=lambda x: x["score"], reverse=True)

    top_candidates = scored[:3]
    selected = top_candidates[0] if top_candidates and top_candidates[0]["score"] > 0 else None

    if not selected:
        return {
            "selected_domain": "Not found in catalog",
            "data_owner": "-",
            "related_sub_team": "-",
            "top_candidates": [],
            "finding": "No confident domain match found in domain catalog."
        }

    d = selected["domain"]

    return {
        "selected_domain": d.get("domain_name", "-"),
        "category": d.get("category", "-"),
        "definition": d.get("definition", "-"),
        "sample_data": d.get("sample_data", "-"),
        "data_owner": d.get("data_owner", "-"),
        "related_sub_team": d.get("related_sub_team", "-"),
        "top_candidates": [
            {
                "domain_name": c["domain"].get("domain_name"),
                "score": c["score"],
                "reasons": c["reasons"]
            }
            for c in top_candidates
        ],
        "finding": f"Selected domain from catalog: {d.get('domain_name')}."
    }


def search_guardrail(user_request, top_k=5):
    chunks = load_json(GUARDRAIL_CHUNKS_PATH)

    request_text = normalize(user_request)
    expanded_request = expand_query(user_request)

    rule_intents = {
        "pii_kvkk": {
            "request_terms": ["pii", "kvkk", "personal data", "kişisel veri", "kisisel veri", "tckn", "phone", "telefon", "msisdn", "location", "lokasyon"],
            "rule_terms": ["kvkk", "kişisel veri", "kisisel veri", "hassas veri", "güvenlik", "guvenlik", "erişim", "erisim", "maskeleme", "anonim"]
        },
        "retention": {
            "request_terms": ["retention", "5 years", "5 year", "historical", "archive", "old data", "eski veri", "geçmiş veri", "gecmis veri"],
            "rule_terms": ["saklama", "retention", "arşiv", "arsiv", "geçmiş", "gecmis", "veri yaşam", "veri yasam", "silme"]
        },
        "approval": {
            "request_terms": ["approval", "exception", "review", "onay", "istisna"],
            "rule_terms": ["onay", "istisna", "veri mimarisi kurulu", "mimari kurul", "review", "uyum", "zorunlu"]
        },
        "platform_architecture": {
            "request_terms": ["source", "system", "platform", "prime", "ods", "ozone", "lakehouse", "kaizen", "dwh"],
            "rule_terms": ["ods", "prime", "dwh", "ozone", "lakehouse", "kaizen", "platform", "mimari", "kaynak sistem"]
        }
    }

    active_intents = []

    for intent, cfg in rule_intents.items():
        if any(normalize(t) in expanded_request for t in cfg["request_terms"]):
            active_intents.append(intent)

    scored = []

    for chunk in chunks:
        raw_text = chunk.get("text", "")
        text = normalize(raw_text)

        score = 0
        reasons = []

        for intent in active_intents:
            for rule_term in rule_intents[intent]["rule_terms"]:
                if normalize(rule_term) in text:
                    score += 10
                    reasons.append(f"{intent}: {rule_term}")

        for term in set(expanded_request.split()):
            if len(term) > 3 and term in text:
                score += 1
                reasons.append(f"keyword: {term}")

        if score > 0:
            enriched_chunk = dict(chunk)
            enriched_chunk["score"] = score
            enriched_chunk["match_reasons"] = reasons[:5]
            scored.append(enriched_chunk)

    scored.sort(key=lambda x: x["score"], reverse=True)

    return scored[:top_k]

def guardian_agent(user_request, request_data):
    guardrail_context = search_guardrail(user_request, top_k=5)

    context_text = "\n\n".join([
        f"Guardrail Chunk {i+1}:\n{c.get('text', '')}"
        for i, c in enumerate(guardrail_context)
    ])

    prompt = f"""
You are the Guardian Agent of Pulse Fabric.

Your job is NOT to copy guardrail text.
Your job is to interpret the retrieved governance guardrails and decide what governance action is required for the incoming data request.
Decision policy:
- If the request only asks to discover which platform, domain, owner or team contains the data, do NOT escalate to Architecture Committee.
- For discovery/routing requests, required approval should be "Data Owner Review / Access Authorization Check".
- Architecture Committee is required only if the request asks for a new platform, new architecture pattern, cross-platform data movement, exception to architecture standards, or non-standard data replication.
- If PII is explicitly requested such as MSISDN, TCKN, phone number, identity number or location history, require KVKK / Privacy Approval + Security Review.
- If historical data or retention period is mentioned, require Retention Policy Review.
- Do not invent approvals that are not required by the request context.
Use only the provided guardrail context and the user request.

Return ONLY valid JSON with this schema:

{{
  "contains_pii": "Yes/No",
  "retention_risk": "Yes/No",
  "risk_level": "LOW/MEDIUM/HIGH",
  "guardian_decision": "Auto Approve/Needs Review/Escalate",
  "required_approval": "- or approval name",
  "guardrail_reasoning": "short business-friendly explanation",
  "applied_guardrail_summary": "short summary of which guardrail logic was applied"
}}

User Request:
{user_request}

Structured Request Data:
{json.dumps(request_data, ensure_ascii=False)}

Retrieved Guardrail Context:
{context_text}
"""

    try:
        response = client.chat.completions.create(
            model=MODEL_NAME,
            messages=[
                {
                    "role": "system",
                    "content": "You are a data governance decision agent. Produce concise, business-friendly governance decisions based on retrieved guardrail context."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.1,
        )

        raw = response.choices[0].message.content
        decision_json = safe_json_loads(raw)

    except Exception as e:
        decision_json = {
            "contains_pii": "Unknown",
            "retention_risk": "Unknown",
            "risk_level": "MEDIUM",
            "guardian_decision": "Needs Review",
            "required_approval": "Manual Governance Review",
            "guardrail_reasoning": f"Guardian Agent could not complete LLM-based guardrail reasoning. Manual review is required. Error: {str(e)}",
            "applied_guardrail_summary": "-"
        }

    return {
        "contains_pii": decision_json.get("contains_pii", "Unknown"),
        "retention_risk": decision_json.get("retention_risk", "Unknown"),
        "architecture_approval_required": decision_json.get("required_approval", "-"),
        "risk_level": decision_json.get("risk_level", "MEDIUM"),
        "guardian_decision": decision_json.get("guardian_decision", "Needs Review"),
        "guardrail_reasoning": decision_json.get("guardrail_reasoning", "-"),
        "required_approval": decision_json.get("required_approval", "-"),
        "matched_guardrail_rules": [],
        "guardrail_sources": [c.get("id", "guardrail") for c in guardrail_context],
        "applied_guardrail_summary": decision_json.get("applied_guardrail_summary", "-")
    }

def get_missing_fields(data):
    return [k for k in REQUIRED_FIELDS if is_empty(data.get(k))]


def inspector_agent(request_data):
    missing = get_missing_fields(request_data)
    maturity_score = int(((len(REQUIRED_FIELDS) - len(missing)) / len(REQUIRED_FIELDS)) * 100)

    return {
        "missing_fields": missing,
        "request_maturity_score": maturity_score,
        "finding": f"{len(missing)} missing field(s) detected."
    }


def strategist_agent(request_data, guardian_result, inspector_result):
    priority = 40
    reasons = []

    if guardian_result.get("risk_level") == "HIGH":
        priority += 35
        reasons.append("High governance risk")

    if guardian_result.get("retention_risk") == "Yes":
        priority += 20
        reasons.append("Retention risk detected")

    if guardian_result.get("contains_pii") == "Yes":
        priority += 20
        reasons.append("PII/KVKK signal detected")

    if inspector_result.get("request_maturity_score", 0) < 70:
        priority -= 10
        reasons.append("Request maturity is not sufficient yet")

    priority = max(0, min(priority, 100))

    if guardian_result.get("guardian_decision") == "Escalate":
        routing = "Route to Guardian approval before delivery assessment"
    elif inspector_result.get("missing_fields"):
        routing = "Collect missing fields before final routing"
    else:
        routing = "Route to responsible data owner and delivery team"

    return {
        "priority_score": priority,
        "routing": routing,
        "reasons": reasons or ["Standard prioritization"]
    }


def scribe_summary(request_data, scout_result, guardian_result, strategist_result):
    return {
        "brd_title": "Pulse Fabric Generated BRD",
        "business_need": request_data.get("business_need", "-"),
        "data_domain": request_data.get("data_domain", "-"),
        "data_owner": scout_result.get("data_owner", "-"),
        "related_sub_team": scout_result.get("related_sub_team", "-"),
        "source_system": request_data.get("source_system", "-"),
        "expected_output": request_data.get("expected_output", "-"),
        "consumer": request_data.get("consumer", "-"),
        "frequency": request_data.get("frequency", "-"),
        "deadline": request_data.get("deadline", "-"),
        "pii_assessment": guardian_result.get("contains_pii", "-"),
        "retention_assessment": guardian_result.get("retention_risk", "-"),
        "risk_level": guardian_result.get("risk_level", "-"),
        "guardian_decision": guardian_result.get("guardian_decision", "-"),
        "priority_score": strategist_result.get("priority_score", "-"),
        "next_step": strategist_result.get("routing", "-"),
    }


def replace_placeholders(doc, values):
    placeholders = {f"{{{{{k}}}}}": str(v) for k, v in values.items()}

    def replace_in_paragraph(paragraph):
        for placeholder, value in placeholders.items():
            if placeholder in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(placeholder, value)

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)


def add_brd_appendix(doc, values):
    doc.add_heading("Pulse Fabric Auto-Filled BRD Summary", level=1)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Value"

    for key, value in values.items():
        row = table.add_row().cells
        row[0].text = key.replace("_", " ").title()
        row[1].text = str(value)


def generate_brd_docx(brd_values):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = os.path.join(OUTPUT_DIR, f"Pulse_Fabric_BRD_{timestamp}.docx")

    if os.path.exists(BRD_TEMPLATE_PATH):
        doc = Document(BRD_TEMPLATE_PATH)
        replace_placeholders(doc, brd_values)
        add_brd_appendix(doc, brd_values)
    else:
        doc = Document()
        doc.add_heading("Pulse Fabric Generated BRD", level=1)
        add_brd_appendix(doc, brd_values)

    doc.save(output_path)
    return output_path


def conductor_agent(request_data, scout_result, guardian_result, inspector_result, strategist_result):
    if inspector_result["missing_fields"]:
        next_step = "Inspector Agent must collect missing fields."
    elif guardian_result.get("guardian_decision") in ["Needs Review", "Escalate"]:
        next_step = "Guardian approval is required before delivery."
    else:
        next_step = "Request is ready for data owner review and delivery assessment."

    return {
        "next_step": next_step,
        "final_status": "Ready for review" if not inspector_result["missing_fields"] else "Needs clarification"
    }


def build_request_data(user_request, scout_result):
    return {
        "business_need": user_request,
        "data_domain": scout_result.get("selected_domain", "-"),
        "source_system": "-",
        "expected_output": "-",
        "consumer": scout_result.get("data_owner", "-"),
        "frequency": "-",
        "contains_pii": "-",
        "deadline": "-",
    }


def run_agents(user_request, existing_data=None):
    scout = scout_agent(user_request)

    request_data = deepcopy(existing_data) if existing_data else build_request_data(user_request, scout)

    if is_empty(request_data.get("data_domain")):
        request_data["data_domain"] = scout.get("selected_domain", "-")

    if is_empty(request_data.get("consumer")):
        request_data["consumer"] = scout.get("data_owner", "-")

    guardian = guardian_agent(user_request, request_data)

    if is_empty(request_data.get("contains_pii")):
        request_data["contains_pii"] = guardian.get("contains_pii", "-")

    inspector = inspector_agent(request_data)
    strategist = strategist_agent(request_data, guardian, inspector)
    conductor = conductor_agent(request_data, scout, guardian, inspector, strategist)

    agent_timeline = {
        "Scout": scout.get("finding", "-"),
        "Inspector": inspector.get("finding", "-"),
        "Guardian": guardian.get("guardrail_reasoning", "-"),
        "Strategist": f"Priority score calculated as {strategist.get('priority_score')}/100.",
        "Scribe": "BRD will be generated after maturity check.",
        "Conductor": conductor.get("next_step", "-"),
    }

    result = {
        "request_data": request_data,
        "scout": scout,
        "guardian": guardian,
        "inspector": inspector,
        "strategist": strategist,
        "conductor": conductor,
        "agent_timeline": agent_timeline,
    }

    return result


def render_table(title, rows):
    md = f"\n## {title}\n\n| Field | Value | Agent |\n|---|---|---|\n"
    for field, value, agent in rows:
        md += f"| **{field}** | {value or '-'} | {agent} |\n"
    return md


def render_result(agent_result, user_request):
    data = agent_result["request_data"]
    scout = agent_result["scout"]
    guardian = agent_result["guardian"]
    inspector = agent_result["inspector"]
    strategist = agent_result["strategist"]
    conductor = agent_result["conductor"]
    timeline = agent_result["agent_timeline"]

    incoming_rows = [
        ("Initial Request", user_request, f"{AGENT_ICONS['Scout']} Scout"),
        ("Business Need", data.get("business_need"), f"{AGENT_ICONS['Scout']} Scout"),
        ("Data Domain", data.get("data_domain"), f"{AGENT_ICONS['Scout']} Scout"),
        ("Data Owner", scout.get("data_owner"), f"{AGENT_ICONS['Scout']} Scout"),
        ("Related Sub-Team", scout.get("related_sub_team"), f"{AGENT_ICONS['Scout']} Scout"),
        ("Source System", data.get("source_system"), f"{AGENT_ICONS['Inspector']} Inspector"),
        ("Expected Output", data.get("expected_output"), f"{AGENT_ICONS['Scribe']} Scribe"),
        ("Consumer", data.get("consumer"), f"{AGENT_ICONS['Conductor']} Conductor"),
        ("Frequency", data.get("frequency"), f"{AGENT_ICONS['Inspector']} Inspector"),
        ("Contains PII", data.get("contains_pii"), f"{AGENT_ICONS['Guardian']} Guardian"),
        ("Deadline", data.get("deadline"), f"{AGENT_ICONS['Strategist']} Strategist"),
    ]

    control_rows = [
        ("Request Maturity Score", f"{inspector.get('request_maturity_score')}/100", f"{AGENT_ICONS['Inspector']} Inspector"),
        ("Priority Score", f"{strategist.get('priority_score')}/100", f"{AGENT_ICONS['Strategist']} Strategist"),
        ("Risk Level", guardian.get("risk_level"), f"{AGENT_ICONS['Guardian']} Guardian"),
        ("Guardian Decision", guardian.get("guardian_decision"), f"{AGENT_ICONS['Guardian']} Guardian"),
        ("Next Step", conductor.get("next_step"), f"{AGENT_ICONS['Conductor']} Conductor"),
    ]

    timeline_rows = [
        (f"{AGENT_ICONS[a]} {a}", timeline.get(a, "-"), "Completed")
        for a in ["Scout", "Inspector", "Guardian", "Strategist", "Scribe", "Conductor"]
    ]

    candidate_lines = []
    for c in scout.get("top_candidates", []):
        candidate_lines.append(f"- **{c['domain_name']}** — score: {c['score']} — {', '.join(c['reasons'])}")

    guardrail_lines = []

    matched_rules = guardian.get("matched_guardrail_rules", [])

    if matched_rules:
        for i, rule in enumerate(matched_rules, 1):
            if isinstance(rule, dict):
                reasons = ", ".join(rule.get("match_reasons", []))
                guardrail_lines.append(
                    f"**Rule {i} — score {rule.get('score', 0)}**  \n"
                    f"{rule.get('rule_text', '-')}  \n"
                    f"_Match reasons: {reasons}_"
                )
            else:
                guardrail_lines.append(f"**Rule {i}:** {rule}")
    else:
        guardrail_lines.append("-")

        matched_rules = guardian.get("matched_guardrail_rules", [])

        if matched_rules:
            for i, rule in enumerate(matched_rules, 1):
                guardrail_lines.append(f"**Rule {i}:** {rule}")
        else:
            for src in guardian.get("guardrail_sources", []):
                guardrail_lines.append(f"- {src}")

    md = ""
    md += render_table("📥 Incoming Request", incoming_rows)
    md += "\n---\n"
    md += render_table("📊 Control Tower", control_rows)
    md += "\n---\n"
    md += render_table("🤖 Agent Timeline", timeline_rows)
    md += "\n---\n"
    md += "## 🦅 Scout Domain Candidates\n"
    md += "\n".join(candidate_lines) if candidate_lines else "-"
    md += "\n\n---\n"
    md += "## 🦏 Guardian Decision Summary\n\n"
    md += f"**Decision:** {guardian.get('guardian_decision', '-')}\n\n"
    md += f"**Risk Level:** {guardian.get('risk_level', '-')}\n\n"
    md += f"**Required Approval:** {guardian.get('required_approval', '-')}\n\n"
    md += f"**Reasoning:** {guardian.get('guardrail_reasoning', '-')}\n"
    md += "\n"

    return md


async def ask_next_missing_field():
    missing_fields = cl.user_session.get("missing_fields") or []
    missing_index = cl.user_session.get("missing_index") or 0

    if missing_index >= len(missing_fields):
        return False

    field = missing_fields[missing_index]

    await cl.Message(content=f"""
## ⚠️ Missing Field {missing_index + 1} of {len(missing_fields)}

🔎 **Inspector Agent**

**{REQUIRED_FIELDS[field]}?**
""").send()

    return True


async def send_brd(agent_result):
    data = agent_result["request_data"]
    scout = agent_result["scout"]
    guardian = agent_result["guardian"]
    strategist = agent_result["strategist"]

    values = scribe_summary(data, scout, guardian, strategist)
    brd_path = generate_brd_docx(values)

    await cl.Message(
        content="🐙 **Scribe Agent** generated the BRD document.",
        elements=[
            cl.File(
                name=os.path.basename(brd_path),
                path=brd_path,
                display="inline"
            )
        ]
    ).send()
    cl.user_session.set("awaiting_next_request_answer", True)

    await cl.Message(content="""
## ✅ Request Completed

Would you like to create another request?

Please type **yes** or **no**.
""").send()
    
async def show_simple_journey_menu():
    cl.user_session.set("user_request", None)
    cl.user_session.set("agent_result", None)
    cl.user_session.set("awaiting_missing_fields", False)
    cl.user_session.set("missing_fields", [])
    cl.user_session.set("missing_index", 0)
    cl.user_session.set("missing_answers", {})
    cl.user_session.set("awaiting_next_request_answer", False)

    await cl.Message(
        content="""
## Select a Customer Journey

| No | Customer Journey |
|---|---|
| **1** | KVKK / PII Risky Request |
| **2** | Expired Retention Data Request |
| **3** | Existing Data Discovery and Routing |

Please type **1**, **2**, **3**, or write your own data request.
"""
    ).send()

async def show_journey_menu():
    cl.user_session.set("user_request", None)
    cl.user_session.set("agent_result", None)
    cl.user_session.set("awaiting_missing_fields", False)
    cl.user_session.set("missing_fields", [])
    cl.user_session.set("missing_index", 0)
    cl.user_session.set("missing_answers", {})
    cl.user_session.set("awaiting_next_request_answer", False)

  
    # if os.path.exists(AGENT_IMAGE_PATH):
     #    elements.append(
      #       cl.Image(
       #          name="Pulse Fabric Digital Teammates",
        #        path=AGENT_IMAGE_PATH,
         #        display="inline"
          #   )
       # )

    await cl.Message(
        content="""
# 🔴 Pulse Fabric

## AI Operating System for Data Lifecycle Management

**Discover • Govern • Protect • Deliver**

Pulse Fabric is an AI-driven control fabric that orchestrates data lifecycle decisions with human governance.

---

## 🎯 Control Tower

| Capability | Target |
|---|---|
| **Current Manual Effort** | 26 MD |
| **Request Maturity Acceleration** | 40–70% |
| **Decision Accuracy Target** | 95%+ |
| **Digital Teammates** | 6 AI Agents |

---

## 🧩 Digital Teammates

| Agent | Role |
|---|---|
| 🦅 **The Scout** | Finds domain, owner and request type |
| 🦉 **The Inspector** | Checks request maturity and missing fields |
| 🐙 **The Scribe** | Generates BRD and analysis drafts |
| 🐺 **The Strategist** | Calculates priority and business impact |
| 🦏 **The Guardian** | Checks KVKK, security and approvals |
| 🐬 **The Conductor** | Creates the final decision package |

---

## 🔄 Lifecycle Flow

**Request → Discover → Govern → Protect → Human Approval → Deliver**

---

## Select a Customer Journey

| No | Customer Journey |
|---|---|
| **1** | KVKK / PII Risky Request |
| **2** | Expired Retention Data Request |
| **3** | Existing Data Discovery and Routing |

Please type **1**, **2**, **3**, or write your own data request.
""",

    ).send()

    


@cl.on_chat_start
async def start():
    await show_journey_menu()


@cl.on_message
async def main(message: cl.Message):
    user_input = message.content.strip()
    awaiting_next = cl.user_session.get("awaiting_next_request_answer")

    if awaiting_next:
        answer = user_input.lower()

        if answer in ["yes", "y", "evet", "e"]:
            await show_simple_journey_menu()
            return

        if answer in ["no", "n", "hayır", "hayir", "h"]:
            cl.user_session.set("awaiting_next_request_answer", False)
            await cl.Message(content="""
✅ Demo flow completed.

Pulse Fabric is ready for the next governance scenario whenever needed.
""").send()
            return

        await cl.Message(content="Please type **yes** or **no**.").send()
        return

    awaiting_missing = cl.user_session.get("awaiting_missing_fields")
    user_request = cl.user_session.get("user_request")
    agent_result = cl.user_session.get("agent_result")

    if awaiting_missing:
        missing_fields = cl.user_session.get("missing_fields") or []
        missing_index = cl.user_session.get("missing_index") or 0
        missing_answers = cl.user_session.get("missing_answers") or {}

        current_field = missing_fields[missing_index]
        missing_answers[current_field] = user_input

        cl.user_session.set("missing_answers", missing_answers)

        next_index = missing_index + 1
        cl.user_session.set("missing_index", next_index)

        if next_index < len(missing_fields):
            await ask_next_missing_field()
            return

        await cl.Message(content="🔄 All missing fields collected. Re-running agents...").send()

        request_data = agent_result["request_data"]
        for key, value in missing_answers.items():
            request_data[key] = value

        final_result = run_agents(user_request, existing_data=request_data)

        cl.user_session.set("agent_result", final_result)
        cl.user_session.set("awaiting_missing_fields", False)
        cl.user_session.set("missing_fields", [])
        cl.user_session.set("missing_index", 0)
        cl.user_session.set("missing_answers", {})

        await cl.Message(content=render_result(final_result, user_request)).send()

        if not final_result["inspector"]["missing_fields"]:
            await send_brd(final_result)

        return

    if user_input in CUSTOMER_JOURNEYS:
        journey = CUSTOMER_JOURNEYS[user_input]
        user_request = journey["sample"]
        journey_name = journey["name"]
    else:
        user_request = user_input
        journey_name = "Custom Request"

    cl.user_session.set("user_request", user_request)

    await cl.Message(content=f"""
## Selected Request

**Journey:** {journey_name}

**Description:**  
{user_request}
""").send()

    await cl.Message(content="🦅 Scout Agent is analyzing the request and searching the domain catalog...").send()

    agent_result = run_agents(user_request)
    cl.user_session.set("agent_result", agent_result)

    await cl.Message(content=render_result(agent_result, user_request)).send()

    missing = agent_result["inspector"]["missing_fields"]

    if missing:
        cl.user_session.set("awaiting_missing_fields", True)
        cl.user_session.set("missing_fields", missing)
        cl.user_session.set("missing_index", 0)
        cl.user_session.set("missing_answers", {})

        await cl.Message(content=f"""
## ⚠️ Missing Fields Detected

**{len(missing)} missing field(s)** found.

Inspector Agent will ask them one by one.  
The request maturity will be recalculated only after all answers are collected.
""").send()

        await ask_next_missing_field()
    else:
        await cl.Message(content="✅ Request is mature. Generating BRD...").send()
        await send_brd(agent_result)